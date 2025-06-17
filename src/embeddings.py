import os
import json
from openai import OpenAI
import pdfplumber
import docx
import time
from datetime import datetime
from typing import Dict, List, Tuple, Optional
import concurrent.futures
from pathlib import Path

EMBEDDINGS_FILE = "embeddings.json"
ERROR_LOG_FILE = "embedding_errors.log"
SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc'}


class FileStatus:
    """Track file processing status"""

    def __init__(self):
        self.valid_files = []
        self.invalid_files = []
        self.skipped_files = []  # Already processed
        self.processing_errors = []
        self.start_time = time.time()


def get_mtime(path):
    """Get file modification time"""
    return os.path.getmtime(path)


def load_embeddings():
    """Load existing embeddings from file"""
    if os.path.exists(EMBEDDINGS_FILE):
        with open(EMBEDDINGS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_embeddings(embeddings):
    """Save embeddings to file"""
    with open(EMBEDDINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(embeddings, f, ensure_ascii=False, indent=2)


def clean_old_embeddings(embeddings: dict, folder: str) -> dict:
    """Remove embeddings for files that no longer exist"""
    folder_key = os.path.basename(folder.rstrip('/'))
    cleaned = {}
    removed_count = 0

    for key, value in embeddings.items():
        if key.startswith(folder_key + '|'):
            filename = key.split('|', 1)[1]
            filepath = os.path.join(folder, filename)
            if os.path.exists(filepath):
                cleaned[key] = value
            else:
                removed_count += 1
        else:
            cleaned[key] = value

    if removed_count > 0:
        print(f"🧹 Cleaned {removed_count} old embeddings from {folder_key}")

    return cleaned


def validate_file(filepath: str) -> Tuple[bool, str]:
    """
    Validate if file can be processed
    Returns: (is_valid, error_message)
    """
    if not os.path.exists(filepath):
        return False, "File does not exist"

    if not os.path.isfile(filepath):
        return False, "Not a file"

    # Check file size (skip if too large - over 10MB)
    file_size = os.path.getsize(filepath)
    if file_size > 10 * 1024 * 1024:
        return False, f"File too large ({file_size / 1024 / 1024:.1f}MB)"

    if file_size == 0:
        return False, "Empty file"

    # Check extension
    ext = Path(filepath).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return False, f"Unsupported file type ({ext})"

    # Try to actually read the file
    try:
        if ext == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read(100)  # Just read first 100 chars
                if not content.strip():
                    return False, "File appears to be empty"

        elif ext == '.pdf':
            with pdfplumber.open(filepath) as pdf:
                if len(pdf.pages) == 0:
                    return False, "PDF has no pages"
                # Try to extract text from first page
                first_page = pdf.pages[0]
                text = first_page.extract_text()
                if not text or not text.strip():
                    return False, "PDF appears to be empty or image-only"

        elif ext in ['.docx', '.doc']:
            doc = docx.Document(filepath)
            if len(doc.paragraphs) == 0:
                return False, "Document has no paragraphs"
            # Check if document has any text
            has_text = any(para.text.strip() for para in doc.paragraphs[:5])
            if not has_text:
                return False, "Document appears to be empty"

    except Exception as e:
        return False, f"Read error: {str(e)}"

    return True, "OK"


def scan_folder(folder: str, existing_embeddings: dict) -> FileStatus:
    """
    Pre-scan folder to categorize files before processing
    """
    status = FileStatus()
    folder_key = os.path.basename(folder.rstrip('/'))

    print(f"\n📂 Scanning {folder}...")

    try:
        files = [f for f in os.listdir(folder) if not f.startswith('.')]
    except Exception as e:
        print(f"❌ Error accessing folder: {e}")
        return status

    for filename in files:
        filepath = os.path.join(folder, filename)
        key = f"{folder_key}|{filename}"

        # Check if already processed
        if key in existing_embeddings:
            mtime = get_mtime(filepath)
            if existing_embeddings[key].get("mtime") == mtime:
                status.skipped_files.append((filename, "Already processed"))
                continue

        # Validate file
        is_valid, error_msg = validate_file(filepath)
        if is_valid:
            status.valid_files.append(filename)
        else:
            status.invalid_files.append((filename, error_msg))

    return status


def load_file_content(filepath: str) -> Optional[str]:
    """Load content from a file"""
    ext = Path(filepath).suffix.lower()

    try:
        if ext == '.txt':
            with open(filepath, encoding='utf-8') as f:
                return f.read()

        elif ext == '.pdf':
            text = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text

        elif ext in ['.docx', '.doc']:
            doc = docx.Document(filepath)
            return "\n".join([para.text for para in doc.paragraphs])

    except Exception as e:
        print(f"Error loading {filepath}: {e}")
        return None

    return None


def create_embedding_for_file(client: OpenAI, filepath: str, key: str, mtime: float) -> Optional[dict]:
    """Create embedding for a single file"""
    text = load_file_content(filepath)
    if not text or not text.strip():
        return None

    try:
        # Limit text size to avoid API errors
        text_for_embedding = text[:8000]

        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=text_for_embedding
        )

        return {
            "embedding": response.data[0].embedding,
            "mtime": mtime,
            "text": text[:4000],  # Store preview for matching
            "processed_at": datetime.now().isoformat()
        }

    except Exception as e:
        print(f"❌ Embedding error for {os.path.basename(filepath)}: {e}")
        return None


def update_embeddings(folder: str, client: OpenAI, progress_callback=None) -> dict:
    """
    Update embeddings with pre-scanning and better error handling
    """
    # Load existing embeddings
    all_embeddings = load_embeddings()

    # Clean old embeddings
    all_embeddings = clean_old_embeddings(all_embeddings, folder)

    # Pre-scan folder
    status = scan_folder(folder, all_embeddings)

    # Display scan results
    print(f"\n📊 Scan Results for {folder}:")
    print(f"✅ Valid files to process: {len(status.valid_files)}")
    print(f"⏭️  Already processed: {len(status.skipped_files)}")
    print(f"❌ Invalid files: {len(status.invalid_files)}")

    if status.invalid_files:
        print("\n⚠️ Invalid files details:")
        for filename, reason in status.invalid_files[:10]:  # Show max 10
            print(f"  - {filename}: {reason}")
        if len(status.invalid_files) > 10:
            print(f"  ... and {len(status.invalid_files) - 10} more")

    # Process valid files
    folder_key = os.path.basename(folder.rstrip('/'))
    updated = False

    for i, filename in enumerate(status.valid_files):
        if progress_callback:
            progress_callback(f"Processing {filename} ({i + 1}/{len(status.valid_files)})")

        filepath = os.path.join(folder, filename)
        key = f"{folder_key}|{filename}"
        mtime = get_mtime(filepath)

        # Add small delay to avoid rate limiting
        if i > 0:
            time.sleep(0.1)

        # Create embedding
        embedding_data = create_embedding_for_file(client, filepath, key, mtime)

        if embedding_data:
            all_embeddings[key] = embedding_data
            updated = True
            print(f"✅ Created embedding for: {filename}")
        else:
            status.processing_errors.append((filename, "Failed to create embedding"))

    # Save if updated
    if updated:
        save_embeddings(all_embeddings)

    # Generate summary
    elapsed_time = time.time() - status.start_time
    print(f"\n📋 Processing Summary for {folder}:")
    print(f"⏱️  Time taken: {elapsed_time:.1f} seconds")
    print(f"✅ Successfully processed: {len(status.valid_files) - len(status.processing_errors)}")
    print(f"❌ Failed: {len(status.processing_errors)}")
    print(f"⏭️  Skipped (already done): {len(status.skipped_files)}")

    # Return only current folder's embeddings
    result = {}
    for k, v in all_embeddings.items():
        if k.startswith(folder_key + '|'):
            result[k] = v

    return result


def get_processing_summary() -> Optional[str]:
    """Get a formatted summary of all processing issues"""
    summary_lines = []

    # Read from error log if exists
    if os.path.exists(ERROR_LOG_FILE):
        with open(ERROR_LOG_FILE, 'r', encoding='utf-8') as f:
            summary_lines = f.readlines()

    if summary_lines:
        return "📝 Processing Summary:\n" + "".join(summary_lines)
    return None


# Backward compatibility functions (from old version)
def log_error(message):
    """Log errors to file for summary at the end"""
    with open(ERROR_LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"{message}\n")


def clear_error_log():
    """Clear error log at start of new run"""
    if os.path.exists(ERROR_LOG_FILE):
        os.remove(ERROR_LOG_FILE)


def get_error_summary():
    """Get summary of all errors - for backward compatibility"""
    return get_processing_summary()


def load_any_file(path):
    """Load any file - for backward compatibility"""
    return load_file_content(path)